#!/usr/bin/env python
"""Apply the post-export corrections to the manuscript and supplementary .docx.

Rewrites only whole paragraphs that are matched by a distinctive anchor string,
preserving each paragraph's style by writing the new text into its first run and
clearing the rest. Every replacement is reported; a paragraph whose anchor is not
found (or is found more than once) is left untouched and flagged, so a
half-applied edit is visible rather than silent.

    python scripts/paper_stats/patch_docx.py [--dry-run]

Back up the .docx files before running (the caller does this).
"""
from __future__ import annotations

import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[2]
DRY = "--dry-run" in sys.argv

MANUSCRIPT = "SV2A_Manuscript_NatNeurosci_draft.docx"
SUPPLEMENT = "SV2A_Supplementary_Information_draft.docx"

# (anchor substring, replacement paragraph text)
MANUSCRIPT_EDITS = [
    (
        "The identical model, thresholds and post-processing were applied",
        "The identical model, thresholds and post-processing were applied to every "
        "recording of both treatment groups, and all detections were reviewed by an "
        "investigator blinded to treatment group; no detector-confidence threshold "
        "was applied. Three animals were excluded from all analyses because of "
        "persistent recording noise that precluded reliable scoring. In total, "
        "1,864 EDF recordings comprising 17,966 channel-hours from 20 chronically "
        "epileptic mice (7 AAV-mDlx-EGFP, 13 AAV-mDlx-SV2A) were analysed, yielding "
        "17,346 seizures (5,288 convulsive, 12,058 non-convulsive) and 2,480,248 "
        "interictal spikes. Batch detection was run on the LUNARC COSMOS cluster "
        "(Lund University). Full details of the architecture, the annotated training "
        "corpus, the training procedure and detection performance are given in the "
        "Supplementary Methods."
    ),
    (
        "Statistical analysis was performed in GraphPad Prism",
        "Statistical analysis was performed in GraphPad Prism [confirm version]. For "
        "all electrographic measures the animal was the unit of analysis (n = 7 "
        "AAV-mDlx-EGFP, n = 13 AAV-mDlx-SV2A), and the baseline and levetiracetam "
        "phases correspond to recording weeks 1–3 and 4–6, respectively. Three "
        "animals were excluded from all analyses because of persistent recording "
        "noise."
    ),
    (
        "Seizure frequency was computed for each animal",
        "Seizure frequency was computed for each animal as the number of detected "
        "events divided by that animal's hours of valid recording in the phase, "
        "separately for convulsive and non-convulsive seizures. The proportion of "
        "convulsive seizure-free days was the percentage of an animal's recorded "
        "days containing no convulsive seizure. Seizure duration was the mean event "
        "duration per animal, computed only for animals in which at least one event "
        "of that class was detected. Interictal spike frequency was the number of "
        "detected spikes divided by the hours of valid recording for that animal and "
        "phase, and spike duration the mean spike duration per animal. Inter-spike "
        "intervals were pooled across all spikes within a group for the "
        "probability-density and cumulative distributions."
    ),
]

SUPPLEMENT_EDITS = [
    (
        "so that no channel contributed windows to both the training",
        "Data were split at the level of the animal, so that no channel contributed "
        "windows to both the training and the validation set and no recording could "
        "leak across the split. The split was stratified by convulsive status: "
        "animals with convulsive events formed one stratum and the remainder "
        "another, with approximately 20% of each stratum assigned to validation, so "
        "that convulsive events were present on both sides of the split rather than "
        "left to a random draw. One channel was excluded from the training corpus "
        "because of persistent recording noise; that animal and two others were also "
        "excluded from all analyses in this paper for the same reason. Because the "
        "corpus derives from a single experimental series, the validation set is an "
        "internal held-out set of animals rather than an independent cohort; this is "
        "discussed in section 8."
    ),
    (
        "Batch detection over the full dataset was run on CPU nodes",
        "Batch detection over the full dataset was run on CPU nodes of the LUNARC "
        "COSMOS cluster, writing results to a SQLite database. All detections were "
        "subsequently reviewed in the NED-Net interface by an investigator blinded "
        "to treatment group, and events judged artefactual were excluded before "
        "analysis; no threshold was applied to the detector's confidence score, "
        "because that score does not cleanly separate true from false detections "
        "(Supplementary Fig. 3e). The identical model, thresholds, post-processing "
        "and review procedure were applied to both treatment groups."
    ),
    (
        "Seizure detector, internal validation. On the held-out validation animals",
        "Seizure detector, internal validation. On the held-out validation animals "
        "(1,505 windows from three animals) the detector achieved an event-level F1 "
        "of 0.82 at the 0.5 operating point (event precision 0.74, event recall "
        "0.91), where an event counted as detected when the predicted and annotated "
        "segments overlapped by more than 20%. Sweeping the threshold gave a maximum "
        "event F1 of 0.84 at a threshold of 0.75 (precision 0.87, recall 0.81; "
        "Supplementary Fig. 3f). The operating point of 0.5 was retained for the "
        "study because it favours sensitivity, the remaining false positives being "
        "removed at the blinded review stage. Per-sample metrics at 0.5 were "
        "precision 0.93 and recall 0.60, the low sample recall reflecting "
        "conservative predicted boundaries rather than missed events; this was the "
        "observation that motivated the boundary hysteresis described in section 7. "
        "These figures come from re-scoring the saved checkpoint against the current "
        "annotation corpus, which is marginally larger than the one present when the "
        "model was trained; the training-time validation figures were F1 0.78 at 0.5 "
        "and 0.81 at 0.70."
    ),
    (
        "Blinded spot-check on study recordings. To estimate the precision",
        "Blinded spot-check on study recordings. To estimate the precision of the "
        "complete pipeline as applied here, 30 recordings were selected from the "
        "analysed dataset and every automatic detection in them was reviewed by the "
        "expert scorer. Of 171 adjudicated detections, 160 were confirmed as genuine "
        "seizures and 11 rejected, a precision of 93.6% (Supplementary Fig. 3d). The "
        "detector's own confidence score was lower for rejected than for confirmed "
        "detections (medians 0.60 and 0.79) but the distributions overlapped "
        "substantially (Supplementary Fig. 3e), so confidence alone does not "
        "separate them and no confidence threshold was applied; every detection "
        "entering the analysis was instead reviewed by eye."
    ),
]

# The Supplementary Figure 3 legend, replacing the bracketed placeholder.
SUPP_FIG3_LEGEND = [
    "(a) Representative automatic detection of a convulsive seizure. Top, "
    "hippocampal EEG (acquired at 2 kHz, shown after decimation to 250 Hz); "
    "bottom, the per-sample seizure probability produced by the U-Net over the "
    "same window. Horizontal lines mark the 0.5 detection threshold and the 0.1 "
    "boundary threshold used to grow the event's onset and offset (hysteresis). "
    "Shading marks the accepted event. The convulsive label comes from the "
    "second-stage classifier applied to the detected event, not from the "
    "probability trace shown. (b) The same for a non-convulsive seizure; the "
    "second-stage classifier returns a low convulsive probability for the same "
    "detected event. (c) Representative interictal spikes detected by the "
    "rule-based detector (arrowheads) in a 12 s epoch. (d) Precision of the "
    "complete pipeline in the blinded spot-check: 30 recordings drawn from the "
    "analysed dataset in which every automatic detection was re-read by the "
    "expert scorer. (e) Detector confidence for the spot-check detections the "
    "scorer confirmed versus rejected; horizontal bars are medians. Rejected "
    "detections carried lower confidence, but the distributions overlap, so "
    "confidence alone does not separate them and the manual review step is not "
    "redundant. (f) Event-level precision, recall and F1 as a function of the "
    "detection threshold, on validation animals held out from training (1,505 "
    "windows from three animals). The vertical line marks the 0.5 operating point "
    "used throughout the study, which favours recall; residual false positives "
    "were removed at the blinded review stage. (g) Confusion matrix for the "
    "second-stage convulsive classifier on held-out validation animals at its "
    "operating threshold of 0.45."
]


def set_text(par, text: str) -> None:
    """Replace a paragraph's text, keeping its style and first run's formatting."""
    if not par.runs:
        par.add_run(text)
        return
    par.runs[0].text = text
    for run in par.runs[1:]:
        run.text = ""


def apply(path: Path, edits) -> int:
    doc = docx.Document(path)
    failures = 0
    for anchor, new_text in edits:
        hits = [p for p in doc.paragraphs if anchor in p.text]
        if len(hits) != 1:
            print(f"  !! {len(hits)} matches for {anchor[:55]!r} — SKIPPED")
            failures += 1
            continue
        print(f"  ok  {anchor[:55]!r}")
        if not DRY:
            set_text(hits[0], new_text)
    if not DRY:
        doc.save(path)
    return failures


def main() -> None:
    total = 0
    print(f"{MANUSCRIPT}:")
    total += apply(ROOT / MANUSCRIPT, MANUSCRIPT_EDITS)
    print(f"{SUPPLEMENT}:")
    total += apply(ROOT / SUPPLEMENT, SUPPLEMENT_EDITS)

    # Supplementary Fig. 3 legend: replace the bracketed placeholder paragraph.
    doc = docx.Document(ROOT / SUPPLEMENT)
    hits = [p for p in doc.paragraphs
            if "Optional. If included, the panels supportable" in p.text]
    if len(hits) == 1:
        print("  ok  Supplementary Fig. 3 legend placeholder")
        if not DRY:
            set_text(hits[0], SUPP_FIG3_LEGEND[0])
            doc.save(ROOT / SUPPLEMENT)
    else:
        print(f"  !! {len(hits)} matches for the Supp. Fig. 3 placeholder — SKIPPED")
        total += 1

    print("\nDRY RUN — nothing written." if DRY else
          f"\nDone. {total} skipped." if total else "\nDone. All edits applied.")


if __name__ == "__main__":
    main()
