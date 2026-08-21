#!/usr/bin/env python
"""Insert the acute-levetiracetam ephys as Extended Data Fig. 1 and renumber.

The new panel is cited just before Figure 5, which is earlier than any existing
Extended Data citation, so by citation order it becomes Extended Data Fig. 1
and the three existing ones shift up: AD spikes -> 2, object location / novel
object -> 3, Barnes maze -> 4. This renames the PDFs, renumbers every citation
in the manuscript and the Extended Data legends, adds the new legend, and
writes the Results text.

Renumbering runs highest-first through placeholders so 1->2 cannot collide with
the 2 that already exists.

    python scripts/paper_stats/patch_lev_ephys_ed.py --dry-run
    python scripts/paper_stats/patch_lev_ephys_ed.py --go
"""
from __future__ import annotations

import copy
import re
import shutil
import sys
from pathlib import Path

import docx

FOLDER = Path("/Users/marcoledri/Dropbox/Work/Manuscripts and papers/"
              "SV2A paper/SV2A paper")
MS = FOLDER / "SV2A_Manuscript_NatNeurosci_draft.docx"
ED = FOLDER / "SV2A_Extended_Data_draft.docx"

# old -> new, applied highest first via placeholders
SHIFT = [(3, 4), (2, 3), (1, 2)]

NEW_RESULTS = (
    "Because SV2A is the target of levetiracetam 12,24, we asked how the gene "
    "therapy interacts with the drug. Levetiracetam is a negative modulator of "
    "release, so supplementing SV2A in interneurons carries a specific "
    "liability: the therapy installs additional target protein at precisely "
    "the synapses through which it acts, and the drug might therefore be "
    "expected to reverse the effect the vector produces. We tested this "
    "directly in acute slices. Bath application of levetiracetam abolished the "
    "increase in sIPSC frequency seen in AAV-mDlx-SV2A animals (7.47 to 3.03 "
    "Hz, P = 0.031, unpaired Welch's t-test), so that treated and control "
    "cells no longer differed (P = 0.40), whereas control cells were little "
    "affected (3.39 to 2.20 Hz, P = 0.30). sIPSC amplitude was unchanged by "
    "the drug in either group, consistent with a presynaptic interaction at "
    "the shared target rather than a change in postsynaptic responsiveness "
    "(Extended Data Fig. 1). We therefore asked whether this synaptic reversal "
    "was sufficient to abolish the benefit of the therapy in vivo."
)

# appended to the end of the existing levetiracetam Results paragraph
NEW_INTERPRETATION = (
    " Levetiracetam thus reduced seizure burden in control animals while "
    "SV2A-treated animals remained below controls throughout treatment, "
    "indicating that the two interventions are compatible but not additive. "
    "Because the vector restricts expression to interneurons, levetiracetam's "
    "action at excitatory terminals should be unchanged between groups, while "
    "its action at inhibitory terminals is amplified only where the therapy "
    "added target, which is consistent with the sub-additivity observed both "
    "in the slice and in vivo. Two caveats temper this interpretation: "
    "SV2A-treated animals entered the treatment phase with few seizures, so "
    "the persistence of their advantage does not by itself establish that the "
    "enhanced inhibition continued to contribute; and the slice recordings "
    "compare separate cells under acute drug application rather than the same "
    "cells before and after, at concentrations that need not match chronic "
    "oral dosing."
)

NEW_LEGEND = (
    "Extended Data Fig. 1 | Acute levetiracetam reverses the SV2A-mediated "
    "increase in inhibitory transmission. Whole-cell voltage-clamp recordings "
    "of spontaneous IPSCs from dentate gyrus granule cells, with and without "
    "acute bath application of levetiracetam. (a) sIPSC frequency, (b) "
    "amplitude and (c) rise time, shown as per-cell medians for AAV-mDlx-EGFP "
    "and AAV-mDlx-SV2A animals under control conditions and during "
    "levetiracetam. Levetiracetam abolished the increase in sIPSC frequency "
    "seen in SV2A-treated animals, so that the groups no longer differed, "
    "while control cells were little affected and amplitude was unchanged "
    "throughout. n = 6 cells (EGFP) and 6 cells (SV2A) under control "
    "conditions, and 7 and 7 during levetiracetam; levetiracetam and control "
    "recordings were made from different cells, so all comparisons are "
    "unpaired (Welch's t-test). Data are mean ± s.e.m. with individual cells "
    "overlaid. *P < 0.05. [Add the acute levetiracetam slice conditions and "
    "the number of animals contributing these cells.]"
)


def renumber(text: str) -> str:
    for old, _new in SHIFT:
        text = text.replace(f"Extended Data Fig. {old}", f"{old}")
    for old, new in SHIFT:
        text = text.replace(f"{old}", f"Extended Data Fig. {new}")
    return text


def set_text(par, text):
    if not par.runs:
        par.text = text
        return
    par.runs[0].text = text
    for r in par.runs[1:]:
        r.text = ""


def insert_after(par, text):
    """Clone a paragraph's XML so the new one inherits its style."""
    new = copy.deepcopy(par._p)
    par._p.addnext(new)
    np = docx.text.paragraph.Paragraph(new, par._parent)
    set_text(np, text)
    return np


def main() -> int:
    go = "--go" in sys.argv
    actions = []

    # ---- 1. rename the PDFs, highest first, via temporary names ----
    pdfs = {n: FOLDER / f"ExtendedData_Figure {n}.pdf" for n in (1, 2, 3)}
    missing = [n for n, p in pdfs.items() if not p.exists()]
    if missing:
        print(f"!! missing PDFs: {missing} — nothing renamed", file=sys.stderr)
    else:
        actions.append("rename ExtendedData_Figure 3 -> 4, 2 -> 3, 1 -> 2")
        if go:
            tmp = {}
            for n in (1, 2, 3):
                t = FOLDER / f"__tmp_ed{n}.pdf"
                pdfs[n].rename(t)
                tmp[n] = t
            for old, new in SHIFT:
                tmp[old].rename(FOLDER / f"ExtendedData_Figure {new}.pdf")

    # ---- 2. renumber citations in both documents ----
    for path in (MS, ED):
        d = docx.Document(str(path))
        n = 0
        for p in d.paragraphs:
            if "Extended Data Fig." in p.text:
                new = renumber(p.text)
                if new != p.text:
                    n += 1
                    if go:
                        set_text(p, new)
        actions.append(f"renumber {n} paragraph(s) in {path.name}")
        if go:
            bak = path.with_suffix(".docx.bak2")
            if not bak.exists():
                shutil.copy2(path, bak)
            d.save(str(path))

    # ---- 3. manuscript: new Results text ----
    d = docx.Document(str(MS))
    hits = [p for p in d.paragraphs
            if p.text.startswith("Because SV2A is the target of levetiracetam")]
    if len(hits) == 1:
        par = hits[0]
        old_first = ("Because SV2A is the target of levetiracetam 12,24, we asked "
                     "how the gene therapy interacts with the drug.")
        # marker unique to the NEW text — the first sentence is shared
        if "Bath application of levetiracetam" in par.text:
            actions.append("Results text already present, skipped")
        elif old_first in par.text:
            actions.append("insert hypothesis + ephys text and interpretation "
                           "into the levetiracetam Results paragraph")
            if go:
                set_text(par, par.text.replace(old_first, NEW_RESULTS)
                         + NEW_INTERPRETATION)
        else:
            actions.append("!! levetiracetam paragraph found but opening "
                           "sentence did not match — NOT edited")
    else:
        actions.append(f"!! levetiracetam Results paragraph: {len(hits)} matches "
                       "— NOT edited")
    if go:
        d.save(str(MS))

    # ---- 4. Extended Data: add the new legend as Fig. 1 ----
    d = docx.Document(str(ED))
    if any(p.text.startswith("Extended Data Fig. 1 | Acute levetiracetam")
           for p in d.paragraphs):
        actions.append("new Extended Data legend already present, skipped")
    else:
        anchor = [p for p in d.paragraphs
                  if p.text.startswith("Extended Data Fig. 2 |")]
        if len(anchor) == 1:
            actions.append("insert the new Extended Data Fig. 1 legend")
            if go:
                prev = anchor[0]._p.getprevious()
                target = (docx.text.paragraph.Paragraph(prev, anchor[0]._parent)
                          if prev is not None else anchor[0])
                new = copy.deepcopy(anchor[0]._p)
                anchor[0]._p.addprevious(new)
                set_text(docx.text.paragraph.Paragraph(new, anchor[0]._parent),
                         NEW_LEGEND)
        else:
            actions.append(f"!! anchor legend 'Extended Data Fig. 2 |': "
                           f"{len(anchor)} matches — legend NOT inserted")
    if go:
        d.save(str(ED))

    for a in actions:
        print(("  " if not a.startswith("!!") else "  ") + a)
    print(f"\n{'APPLIED' if go else 'DRY RUN — re-run with --go'}")
    return 1 if any(a.startswith("!!") for a in actions) else 0


if __name__ == "__main__":
    raise SystemExit(main())
