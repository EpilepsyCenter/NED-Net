#!/usr/bin/env python
"""Add Kolmogorov-Smirnov D values to the figure legends in the manuscript.

Every cumulative-distribution panel in Figures 2, 4 and 5 is compared with the
same two-sample K-S test, so marking some panels and not others leaves a reader
unable to tell whether an unmarked panel failed the test or was never tested.
D is added for all of them: it shows the gradient the asterisks flatten
(D = 0.24 for the clear effects against 0.09 for the marginal ones).

The K-S comparisons pool events within a group and treat each event as
independent, when the real unit is the cell (Figure 2) or the animal
(Figures 4, 5); at n = 600-10^6 events that makes P values very small for
slight differences. The legends now say so, and point to the cell- and
animal-level tests as the inferential ones.

Also corrects the Figure 2 Results sentence: mean amplitude and rise time do
not differ, but the corresponding distributions do, which the previous wording
denied.

    python scripts/paper_stats/patch_ks_legends.py --dry-run
    python scripts/paper_stats/patch_ks_legends.py --go
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import docx

DOC = Path("/Users/marcoledri/Dropbox/Work/Manuscripts and papers/"
           "SV2A paper/SV2A paper/SV2A_Manuscript_NatNeurosci_draft.docx")

# (anchor found in the paragraph, old substring, new substring)
EDITS = [
    # --- Figure 2 legend: D for all six cumulative panels + pooling caveat
    ("(d,e) Cumulative inter-event-interval distributions",
     "cumulative inter-event-interval distributions by two-sample "
     "Kolmogorov–Smirnov test. *P < 0.05.",
     "cumulative distributions by two-sample Kolmogorov–Smirnov test "
     "(D = 0.225 in d, 0.147 in e, 0.105 in g, 0.089 in h, 0.237 in j and "
     "0.139 in k). Kolmogorov–Smirnov comparisons were made on events "
     "pooled within group and are descriptive; the cell-level comparisons in "
     "c, f and i provide the inferential statistics. *P < 0.05."),

    # --- Figure 4 legend: D for panel f
    ("(e) Probability-density and (f) cumulative distributions",
     "inter-spike-interval distributions compared by two-sample "
     "Kolmogorov–Smirnov test.",
     "inter-spike-interval distributions compared by two-sample "
     "Kolmogorov–Smirnov test on intervals pooled within group "
     "(f, D = 0.125)."),

    # --- Figure 5 legend: D for both curves in panel h
    ("(h) Inter-spike-interval distributions and (i) spike duration.",
     "(h) Inter-spike-interval distributions and (i) spike duration.",
     "(h) Inter-spike-interval distributions, compared by two-sample "
     "Kolmogorov–Smirnov test on intervals pooled within group "
     "(EGFP D = 0.104, SV2A D = 0.066), and (i) spike duration."),

    # --- Methods: state that K-S comparisons pool events
    ("Inter-spike-interval distributions were compared",
     "Inter-spike-interval distributions were compared with the two-sample "
     "Kolmogorov–Smirnov test.",
     "Inter-spike-interval distributions were compared with the two-sample "
     "Kolmogorov–Smirnov test on intervals pooled within group. Because "
     "pooling treats each interval as independent when the unit of analysis "
     "is the animal, these comparisons are descriptive and the D statistic "
     "is reported with the P value."),

    ("For electrophysiological data, mean differences",
     "Distributions were analyzed with two-sample Kolmogorov–Smirnov test.",
     "Distributions were analyzed with the two-sample Kolmogorov–Smirnov "
     "test on events pooled within group; as above, these comparisons are "
     "descriptive relative to the cell-level tests, and D is reported in the "
     "figure legend."),

    # --- Results: the amplitude distributions cross rather than shift
    ("The corresponding cumulative distributions differed modestly",
     "The corresponding cumulative distributions differed modestly for "
     "amplitude (Fig. 2g,h) and, more substantially, for sIPSC rise time, "
     "which was right-shifted in SV2A-treated animals, indicating "
     "slower-rising events (Fig. 2j,k, Kolmogorov–Smirnov test).",
     "The corresponding cumulative distributions differed modestly for "
     "amplitude (Fig. 2g,h), but these curves cross rather than separate: "
     "median amplitude was unchanged (sIPSC 40.3 versus 40.0 pA) while the "
     "SV2A distributions were narrower at both extremes, with fewer very "
     "small and fewer very large events, so the difference reflects reduced "
     "spread rather than a shift in either direction. The sIPSC rise-time "
     "distribution differed more substantially and in a single direction, "
     "being right-shifted in SV2A-treated animals and indicating "
     "slower-rising events (Fig. 2j,k, Kolmogorov–Smirnov test)."),

    # --- Methods: flag the acute-LEV slice conditions for the postdoc
    ("Whole-cell voltage-clamp recordings of spontaneous and miniature IPSCs",
     "[Slice preparation, solutions, recording and analysis details to be "
     "added from the electrophysiology protocol.]",
     "A separate set of recordings was made with levetiracetam applied "
     "acutely to the slice. [Slice preparation, solutions, recording and "
     "analysis details to be added from the electrophysiology protocol. "
     "For the acute levetiracetam experiments, add: concentration, carrier/"
     "vehicle, pre-incubation and bath-application times, whether "
     "levetiracetam and control recordings came from the same slices and "
     "animals or from separate ones, and the number of animals contributing "
     "the 7 EGFP and 7 SV2A cells.]"),

    # --- Results text: means vs distributions
    ("Cumulative distributions confirmed a leftward shift",
     "Neither the amplitude nor the rise time of sIPSCs or mIPSCs differed "
     "between groups (Fig. 2f–k). The cumulative distribution of sIPSC "
     "rise times was, however, right-shifted in SV2A-treated animals, "
     "indicating slower-rising events (Fig. 2i–k, "
     "Kolmogorov–Smirnov test).",
     "Mean amplitude and rise time did not differ between groups for either "
     "sIPSCs or mIPSCs (Fig. 2f,i). The corresponding cumulative "
     "distributions differed modestly for amplitude (Fig. 2g,h) and, more "
     "substantially, for sIPSC rise time, which was right-shifted in "
     "SV2A-treated animals, indicating slower-rising events "
     "(Fig. 2j,k, Kolmogorov–Smirnov test)."),
]


def set_text(par, text):
    """Replace a paragraph's text, keeping its style via the first run."""
    if not par.runs:
        par.text = text
        return
    par.runs[0].text = text
    for r in par.runs[1:]:
        r.text = ""


def main() -> int:
    go = "--go" in sys.argv
    if not DOC.exists():
        print(f"not found: {DOC}", file=sys.stderr)
        return 1
    d = docx.Document(str(DOC))

    applied, failed = 0, []
    for anchor, old, new in EDITS:
        if any(new in p.text for p in d.paragraphs):
            print(f"\n--- {anchor[:60]}\n  (already applied, skipped)")
            applied += 1
            continue
        hits = [p for p in d.paragraphs if anchor in p.text and old in p.text]
        if len(hits) != 1:
            failed.append((anchor[:52], f"{len(hits)} paragraphs matched"))
            continue
        par = hits[0]
        print(f"\n--- {anchor[:60]}")
        print(f"  OLD: ...{old[:150]}")
        print(f"  NEW: ...{new[:150]}")
        if go:
            set_text(par, par.text.replace(old, new))
        applied += 1

    print(f"\n{applied}/{len(EDITS)} edits {'applied' if go else 'ready'}")
    for a, why in failed:
        print(f"  !! NOT APPLIED: {a} ({why})")
    if failed:
        print("  (a failed anchor leaves that paragraph untouched, so a "
              "half-applied edit is visible rather than silent)")

    if go and applied:
        bak = DOC.with_suffix(".docx.bak")
        if not bak.exists():
            shutil.copy2(DOC, bak)
            print(f"\nbacked up -> {bak.name}")
        d.save(str(DOC))
        print(f"saved {DOC.name}")
    elif not go:
        print("\nDry run. Re-run with --go to write.")
    return 0 if not failed else 1


if __name__ == "__main__":
    raise SystemExit(main())
